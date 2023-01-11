import datetime
import os
import random
import smtplib
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.utils import formataddr, parseaddr

from django.utils import timezone
from docx import Document

from file.models import EmailModel
from file.utils import getJsonConfig


def sendEmail(email_obj: EmailModel):
    '''这个函数只应该被用来送出激活邮箱的验证码'''

    # https://blog.csdn.net/qq_41187577/article/details/119004959
    def _format_addr(s):
        addr = parseaddr(s)
        return formataddr(addr)

    try:
        if (timezone.now() - email_obj.last_send_code_time).seconds <= 600:
            return
    except:
        pass

    email_obj.code = random.randint(1, 999999999)
    email_obj.save()
    fromaddr = getJsonConfig('do_push_email')
    pwd = getJsonConfig('do_push_email_pwd')
    smtp_url = getJsonConfig('do_push_email_smtp_url')
    toaddrs = [email_obj.email]

    # 写docx
    file_name = 'KindleLN验证码：[%i]_for_%s.docx' % (
        email_obj.code,
        email_obj.email.replace('@kindle.com', '_kindle_com'),
    )
    doc = Document()
    doc.add_paragraph('您好，%s！' % email_obj.email)
    doc.add_paragraph('这是您在KindleLN的验证码：')
    code_para = doc.add_paragraph()
    code_para.add_run('%i' % email_obj.code).blod = True
    doc.add_paragraph('')
    doc.add_paragraph(
        '邮件由系统自动在【%s】发送，有效期10分钟。'
        % datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    )
    doc.save(file_name)

    # 发送邮件
    content = '[%s]UTC+8\nTO:%s\nCODE:%i' % (
        datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        str(toaddrs),
        email_obj.code,
    )
    text_apart = MIMEText(content)

    file_apart = MIMEApplication(open(file_name, 'rb').read())
    file_apart.add_header('Content-Disposition', 'attachment', filename=file_name)

    mail = MIMEMultipart()
    mail.attach(text_apart)
    mail.attach(file_apart)
    mail['Subject'] = 'code_for_%s' % email_obj.email.replace(
        '@kindle.com', '_kindle_com'
    )
    # 自定义发件人和收件人信息
    # （可以在收到邮件的时候看到发件人和收件人信息），没设置该部分在收邮件时告知是未知发件人
    # 所以Kindle会当没看到，乐
    # 自定义发件人名称
    mail['From'] = _format_addr('KindleLN <%s>' % fromaddr)

    # 自定义收件人(不定义会显示'收件人未填写')
    for item in toaddrs:
        mail['to'] = _format_addr(item)

    try:
        server = smtplib.SMTP(smtp_url)
        server.login(fromaddr, pwd)
        server.sendmail(fromaddr, toaddrs, mail.as_string())
        server.quit()
    except:
        raise
    finally:
        os.remove(file_name)
